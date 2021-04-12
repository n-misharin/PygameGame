from docx import Document
from docx.shared import Pt


document = Document()

font = document.styles['Normal'].font
font.name = 'Times New Roman'
font.size = Pt(14)

result = []

data = open('data.txt', encoding="utf8").read()
for row in data.split('\n')[1:]:
    if len(row) > 1:
        result.append(row.split('\t'))

p_title = document.add_paragraph('')
p_title.add_run('Содержание учебного плана программы').italic = True


def add_chapter(title=''):
    p = document.add_paragraph('')
    p.add_run(f'Раздел {title}').bold = True


def add_chapter_theme(theme='', theory='', practice=''):
    p = document.add_paragraph('')
    p.add_run(f'{theme}').bold = True
    if theory:
        p2 = document.add_paragraph('')
        p2.add_run('Теория: ').italic = True
        p2.add_run(theory)
    if practice:
        p3 = document.add_paragraph('')
        p3.add_run('Практика: ').italic = True
        p3.add_run(practice)


def is_chapter_name(line):
    return len(line[0]) > 6


for line in result:
    if is_chapter_name(line):
        add_chapter(line[0])
    else:
        practice = line[6] if len(line) >= 7 else ''
        add_chapter_theme(theme=line[0] + '. ' + line[1], theory=line[5], practice=practice)

document.save('test.docx')